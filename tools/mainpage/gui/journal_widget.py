"""
journal_widget — Tagesprotokoll-Widget des Mainpage-Dashboards.

Zeigt Journal-Einträge des aktuellen Tages mit Datums-Navigation
und Möglichkeit zum Anlegen neuer Notizen.

Author: Patrick Riederich
Version: 1.0
"""

from __future__ import annotations

from datetime import date, timedelta

from PySide6.QtCore import Qt
from PySide6.QtWidgets import (
    QDialog,
    QFileDialog,
    QFrame,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QPushButton,
    QScrollArea,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from core import theme
from core.dialogs import FinlaiInfoDialog, FinlaiSuccessDialog
from core.icons import Icons, get_icon
from core.logger import get_logger
from tools.mainpage.application.journal_service import JournalService

_log = get_logger(__name__)

_WEEKDAYS = [
    "Montag",
    "Dienstag",
    "Mittwoch",
    "Donnerstag",
    "Freitag",
    "Samstag",
    "Sonntag",
]
_MONTHS = [
    "Januar",
    "Februar",
    "März",
    "April",
    "Mai",
    "Juni",
    "Juli",
    "August",
    "September",
    "Oktober",
    "November",
    "Dezember",
]

_ENTRY_ICONS = {
    "note": "[Notiz]",
    "task_done": "[OK]",
    "tool_used": "[Tool]",
    "auto": "[Auto]",
}


def _format_date(d: date) -> str:
    return f"{_WEEKDAYS[d.weekday()]}, {d.day}. {_MONTHS[d.month - 1]} {d.year}"


def _note_input_css(t) -> str:
    """Gibt CSS für QLineEdit/QTextEdit-Felder im Notiz-Dialog zurück."""
    return (
        f"background: {t.BG_INPUT}; color: {t.TEXT_MAIN}; "
        f"border: 1px solid {t.BORDER}; border-radius: 6px; padding: 4px 8px;"
        f" font-family: 'JetBrains Mono', 'Consolas', monospace; font-size: 13px;"
    )


class NewNoteDialog(QDialog):
    """Kleiner Dialog zum Anlegen einer manuellen Notiz."""

    def __init__(self, parent: QWidget | None = None) -> None:
        """Initialisiert den Dialog.

        Args:
            parent: Optionales Eltern-Widget.
        """
        super().__init__(parent)
        self.setWindowTitle("Neue Notiz")
        self.setMinimumWidth(420)
        self._build_ui()
        from core import theme as _t  # noqa: PLC0415

        _t.register_listener(self.apply_theme)

    def _build_ui(self) -> None:
        """Erstellt die Dialog-Oberfläche."""
        t = theme.get()
        self.setStyleSheet(
            f"QDialog {{ background: {t.BG_MAIN}; color: {t.TEXT_MAIN}; }}"
        )

        lyt = QVBoxLayout(self)
        lyt.setSpacing(16)
        lyt.setContentsMargins(24, 24, 24, 24)

        def _lbl(text: str) -> QLabel:
            lbl_widget = QLabel(text)
            lbl_widget.setStyleSheet(
                f"font-family: 'Raleway'; font-size: 13px; "
                f"color: {t.TEXT_DIM}; font-weight: bold;"
            )
            return lbl_widget

        lyt.addWidget(_lbl("Titel:"))
        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText("Notiz-Titel …")
        self.title_edit.setFixedHeight(36)
        self.title_edit.setStyleSheet(_note_input_css(t))
        lyt.addWidget(self.title_edit)

        lyt.addWidget(_lbl("Notiz:"))
        self.content_edit = QTextEdit()
        self.content_edit.setFixedHeight(100)
        self.content_edit.setPlaceholderText("Optionaler Inhalt …")
        self.content_edit.setStyleSheet(_note_input_css(t))
        lyt.addWidget(self.content_edit)

        lyt.addSpacing(8)
        btn_row = QHBoxLayout()
        btn_row.addStretch()

        btn_cancel = QPushButton("Abbrechen")
        btn_cancel.setIcon(get_icon(Icons.CLOSE))
        btn_cancel.setFixedHeight(36)
        btn_cancel.setStyleSheet(
            f"QPushButton {{ background: {t.BG_BUTTON}; color: {t.TEXT_MAIN};"
            f" border: 1px solid {t.BORDER}; border-radius: 6px;"
            f" font-family: 'Raleway'; font-size: 13px; padding: 0 16px; }}"
            f"QPushButton:hover {{ background: {t.BG_SIDEBAR_HOVER}; }}"
        )
        btn_cancel.clicked.connect(self.reject)
        btn_row.addWidget(btn_cancel)

        btn_save = QPushButton("Speichern")
        btn_save.setIcon(get_icon(Icons.SAVE))
        btn_save.setFixedHeight(36)
        btn_save.setStyleSheet(
            f"QPushButton {{ background: {t.ACCENT}; color: {t.BG_DARK};"
            f" border: none; border-radius: 6px;"
            f" font-family: 'Raleway'; font-size: 13px; font-weight: bold; padding: 0 16px; }}"
            f"QPushButton:hover {{ background: {t.BG_SIDEBAR_HOVER}; }}"
        )
        btn_save.clicked.connect(self.accept)
        btn_row.addWidget(btn_save)

        lyt.addLayout(btn_row)

    def apply_theme(self) -> None:
        """Aktualisiert Widget-Farben für das aktive Theme."""
        from core import theme  # noqa: PLC0415

        c = theme.get()
        self.setStyleSheet(
            f"QDialog {{ background: {c.BG_MAIN}; color: {c.TEXT_MAIN}; }}"
        )
        self.title_edit.setStyleSheet(_note_input_css(c))
        self.content_edit.setStyleSheet(_note_input_css(c))

    @property
    def note_title(self) -> str:
        return self.title_edit.text().strip()

    @property
    def note_content(self) -> str:
        return self.content_edit.toPlainText().strip()


class JournalWidget(QWidget):
    """Tagesprotokoll-Widget mit Eintragsanzeige und Datums-Navigation.

    Zeigt alle Journal-Einträge des gewählten Tages und erlaubt
    das Anlegen neuer Notizen.
    """

    def __init__(
        self,
        journal: JournalService,
        parent: QWidget | None = None,
    ) -> None:
        """Initialisiert das Journal-Widget.

        Args:
            journal: JournalService-Instanz.
            parent: Optionales Eltern-Widget.
        """
        super().__init__(parent)
        self._journal = journal
        self._current_date = date.today()
        t = theme.get()
        self.setStyleSheet(
            f"background-color: {t.CARD_BG}; "
            f"border: 1px solid {t.BORDER}; border-radius: 4px;"
        )

        outer = QVBoxLayout(self)
        outer.setContentsMargins(12, 8, 12, 8)
        outer.setSpacing(6)

        # Header-Zeile
        hdr = QHBoxLayout()
        self._title_lbl = QLabel("Tagesprotokoll")
        self._title_lbl.setStyleSheet(
            f"font-family: 'Raleway', 'Segoe UI', sans-serif; "
            f"font-size: 13px; font-weight: bold; color: {t.ACCENT}; "
            f"background: transparent; border: none;"
        )
        hdr.addWidget(self._title_lbl)
        hdr.addStretch()

        self._btn_note = QPushButton("+ Notiz")
        self._btn_note.setFixedHeight(24)
        self._btn_note.setStyleSheet(
            f"QPushButton {{ font-size: 11px; padding: 2px 8px; "
            f"background: {t.BG_BUTTON}; color: {t.TEXT_MAIN}; "
            f"border: 1px solid {t.ACCENT}; border-radius: 3px; }}"
            f"QPushButton:hover {{ border: 2px solid {t.ACCENT}; color: {t.ACCENT}; }}"
            f"QPushButton:pressed {{ background: {t.ACCENT}; color: {t.BG_MAIN}; }}"
        )
        self._btn_note.clicked.connect(self._on_add_note)
        hdr.addWidget(self._btn_note)

        self._btn_export = QPushButton()
        self._btn_export.setIcon(get_icon(Icons.EXPORT))
        self._btn_export.setFixedSize(24, 24)
        self._btn_export.setToolTip("Als DOCX exportieren")
        self._btn_export.setStyleSheet(
            f"QPushButton {{ background: {t.BG_BUTTON}; color: {t.TEXT_MAIN}; "
            f"border: 1px solid {t.BORDER}; border-radius: 3px; }}"
            f"QPushButton:hover {{ border-color: {t.ACCENT}; color: {t.ACCENT}; }}"
            f"QPushButton:pressed {{ background: {t.ACCENT}; color: {t.BG_MAIN}; }}"
        )
        self._btn_export.clicked.connect(self._on_export_docx)
        hdr.addWidget(self._btn_export)
        outer.addLayout(hdr)

        # Datum-Anzeige
        self._date_lbl = QLabel()
        self._date_lbl.setStyleSheet(
            f"font-size: 12px; color: {t.TEXT_DIM}; background: transparent; border: none;"
        )
        outer.addWidget(self._date_lbl)

        self._sep1 = QFrame()
        self._sep1.setFrameShape(QFrame.Shape.HLine)
        self._sep1.setFixedHeight(1)
        self._sep1.setStyleSheet(f"background: {t.BORDER}; border: none;")
        outer.addWidget(self._sep1)

        # Eintrags-Liste
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("border: none; background: transparent;")

        self._content = QWidget()
        self._content.setStyleSheet("background: transparent;")
        self._content_lyt = QVBoxLayout(self._content)
        self._content_lyt.setContentsMargins(0, 0, 0, 0)
        self._content_lyt.setSpacing(3)
        self._content_lyt.addStretch()

        scroll.setWidget(self._content)
        outer.addWidget(scroll)

        # Navigation
        self._sep2 = QFrame()
        self._sep2.setFrameShape(QFrame.Shape.HLine)
        self._sep2.setFixedHeight(1)
        self._sep2.setStyleSheet(f"background: {t.BORDER}; border: none;")
        outer.addWidget(self._sep2)

        nav = QHBoxLayout()
        self._btn_prev = QPushButton("< Gestern")
        self._btn_today = QPushButton("Heute")
        self._btn_next = QPushButton("Morgen >")
        self._nav_buttons = (self._btn_prev, self._btn_today, self._btn_next)
        for btn in self._nav_buttons:
            btn.setFixedHeight(22)
            btn.setStyleSheet(
                f"QPushButton {{ font-size: 11px; padding: 2px 8px; "
                f"background: {t.BG_BUTTON}; color: {t.TEXT_MAIN}; "
                f"border: 1px solid {t.BORDER}; border-radius: 3px; }}"
                f"QPushButton:hover {{ border-color: {t.ACCENT}; color: {t.ACCENT}; }}"
                f"QPushButton:pressed {{ background: {t.ACCENT}; color: {t.BG_MAIN}; }}"
            )
        self._btn_prev.clicked.connect(lambda: self._navigate(-1))
        self._btn_today.clicked.connect(self._go_today)
        self._btn_next.clicked.connect(lambda: self._navigate(1))
        nav.addWidget(self._btn_prev)
        nav.addWidget(self._btn_today)
        nav.addWidget(self._btn_next)
        nav.addStretch()
        outer.addLayout(nav)

        self._refresh()
        from core import theme as _t  # noqa: PLC0415

        _t.register_listener(self.apply_theme)

    def apply_theme(self) -> None:
        """Aktualisiert Widget-Farben für das aktive Theme."""
        from core import theme  # noqa: PLC0415

        c = theme.get()
        self.setStyleSheet(
            f"background-color: {c.CARD_BG}; "
            f"border: 1px solid {c.BORDER}; border-radius: 4px;"
        )
        self._title_lbl.setStyleSheet(
            f"font-family: 'Raleway', 'Segoe UI', sans-serif; "
            f"font-size: 13px; font-weight: bold; color: {c.ACCENT}; "
            f"background: transparent; border: none;"
        )
        self._btn_note.setStyleSheet(
            f"QPushButton {{ font-size: 11px; padding: 2px 8px; "
            f"background: {c.BG_BUTTON}; color: {c.TEXT_MAIN}; "
            f"border: 1px solid {c.ACCENT}; border-radius: 3px; }}"
            f"QPushButton:hover {{ border: 2px solid {c.ACCENT}; color: {c.ACCENT}; }}"
            f"QPushButton:pressed {{ background: {c.ACCENT}; color: {c.BG_MAIN}; }}"
        )
        self._date_lbl.setStyleSheet(
            f"font-size: 12px; color: {c.TEXT_DIM}; background: transparent; border: none;"
        )
        self._sep1.setStyleSheet(f"background: {c.BORDER}; border: none;")
        self._sep2.setStyleSheet(f"background: {c.BORDER}; border: none;")
        for btn in self._nav_buttons:
            btn.setStyleSheet(
                f"QPushButton {{ font-size: 11px; padding: 2px 8px; "
                f"background: {c.BG_BUTTON}; color: {c.TEXT_MAIN}; "
                f"border: 1px solid {c.BORDER}; border-radius: 3px; }}"
                f"QPushButton:hover {{ border-color: {c.ACCENT}; color: {c.ACCENT}; }}"
                f"QPushButton:pressed {{ background: {c.ACCENT}; color: {c.BG_MAIN}; }}"
            )
        self._refresh()

    def _refresh(self) -> None:
        """Lädt und zeigt die Einträge des aktuellen Datums."""
        self._date_lbl.setText(_format_date(self._current_date))
        self._btn_next.setEnabled(self._current_date < date.today())

        # Alte Einträge löschen
        while self._content_lyt.count() > 1:
            item = self._content_lyt.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        date_str = self._current_date.isoformat()
        entries = self._journal.get_by_date(date_str)

        t = theme.get()
        if not entries:
            lbl = QLabel("Keine Einträge für diesen Tag.")
            lbl.setStyleSheet(
                f"color: {t.TEXT_DIM}; font-size: 12px; background: transparent; border: none;"
            )
            self._content_lyt.insertWidget(0, lbl)
            return

        for i, entry in enumerate(entries):
            row = self._build_entry_row(entry)
            self._content_lyt.insertWidget(i, row)

    def _build_entry_row(self, entry) -> QWidget:
        """Baut eine Eintrags-Zeile auf.

        Args:
            entry: JournalEntry-Instanz.

        Returns:
            QWidget mit Icon, Zeit und Titel.
        """
        t = theme.get()
        row = QWidget()
        row.setStyleSheet("background: transparent; border: none;")
        lyt = QHBoxLayout(row)
        lyt.setContentsMargins(0, 2, 0, 2)
        lyt.setSpacing(8)

        # Uhrzeit
        ts = entry.timestamp
        try:
            from datetime import datetime

            time_str = datetime.fromisoformat(ts).strftime("%H:%M")
        except (ValueError, TypeError):
            time_str = "—"

        time_lbl = QLabel(time_str)
        time_lbl.setFixedWidth(38)
        time_lbl.setStyleSheet(
            f"font-size: 11px; color: {t.TEXT_DIM}; background: transparent; border: none;"
        )
        lyt.addWidget(time_lbl)

        # Icon
        icon = _ENTRY_ICONS.get(entry.entry_type, "[+]")
        icon_lbl = QLabel(icon)
        icon_lbl.setStyleSheet(
            "font-size: 12px; background: transparent; border: none;"
        )
        lyt.addWidget(icon_lbl)

        # Titel — Plain-Text erzwingen: Journal-Titel enthalten untrusted
        # Task-Titel (Software-Namen aus winget/Registry; R22/-
        # Klasse) — seit auch voll-automatisch via add_task_auto_done.
        title = entry.title
        title_lbl = QLabel(title)
        title_lbl.setTextFormat(Qt.TextFormat.PlainText)
        title_lbl.setStyleSheet(
            f"font-size: 12px; color: {t.TEXT_MAIN}; background: transparent; border: none;"
        )
        lyt.addWidget(title_lbl)
        lyt.addStretch()

        return row

    def _navigate(self, delta: int) -> None:
        """Wechselt den angezeigten Tag.

        Args:
            delta: +1 für morgen, -1 für gestern.
        """
        self._current_date += timedelta(days=delta)
        if self._current_date > date.today():
            self._current_date = date.today()
        self._refresh()

    def _go_today(self) -> None:
        """Springt zum heutigen Tag zurück."""
        self._current_date = date.today()
        self._refresh()

    def _on_export_docx(self) -> None:
        """Exportiert die Einträge des aktuellen Tages als DOCX-Datei."""
        try:
            from docx import Document  # noqa: PLC0415
            from docx.shared import Pt, RGBColor  # noqa: PLC0415
        except ImportError:
            FinlaiInfoDialog(
                title="python-docx fehlt",
                message="Bitte installieren: pip install python-docx",
                icon_name=Icons.WARNING,
                parent=self,
            ).exec()
            return

        date_str = self._current_date.isoformat()
        entries = self._journal.get_by_date(date_str)
        formatted = _format_date(self._current_date)

        default_name = f"Tagesprotokoll_{date_str}.docx"
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Tagesprotokoll exportieren",
            default_name,
            "Word-Dokumente (*.docx)",
        )
        if not path:
            return

        try:
            doc = Document()

            # Titel
            title_para = doc.add_heading(level=1)
            run = title_para.add_run(f"FINLAI Tagesprotokoll — {formatted}")
            run.font.color.rgb = RGBColor(0x00, 0xD4, 0xFF)

            if not entries:
                doc.add_paragraph("Keine Einträge für diesen Tag.")
            else:
                for entry in entries:
                    try:
                        from datetime import datetime as _dt  # noqa: PLC0415

                        time_str = _dt.fromisoformat(entry.timestamp).strftime("%H:%M")
                    except (ValueError, TypeError):
                        time_str = "—"

                    type_label = _ENTRY_ICONS.get(entry.entry_type, "[+]")
                    para = doc.add_paragraph()
                    run_time = para.add_run(f"{time_str}  {type_label}  ")
                    run_time.font.size = Pt(9)
                    run_title = para.add_run(entry.title)
                    run_title.bold = True

                    if entry.content:
                        doc.add_paragraph(
                            entry.content
                        ).paragraph_format.left_indent = Pt(20)

            doc.save(path)
            FinlaiSuccessDialog(
                title="Export erfolgreich",
                message="Tagesprotokoll gespeichert.",
                file_path=str(path),
                parent=self,
            ).exec()
        except Exception as exc:  # noqa: BLE001
            _log.error("DOCX-Export fehlgeschlagen: %s", exc)
            FinlaiInfoDialog(
                title="Export fehlgeschlagen",
                message=f"Fehler beim Erstellen der DOCX-Datei:\n{exc}",
                icon_name=Icons.ERROR,
                parent=self,
            ).exec()

    def _on_add_note(self) -> None:
        """Öffnet Dialog für neue Notiz."""
        dlg = NewNoteDialog(self)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        title = dlg.note_title
        if not title:
            return
        try:
            self._journal.add_note(title=title, content=dlg.note_content)
            if self._current_date == date.today():
                self._refresh()
        except Exception as exc:
            _log.error("Notiz konnte nicht gespeichert werden: %s", exc)

    def refresh(self) -> None:
        """Aktualisiert die Eintrags-Anzeige."""
        self._refresh()
